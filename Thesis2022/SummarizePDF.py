from docx import *
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize,sent_tokenize
from PyPDF2 import PdfFileReader

def summ(self,path,filename):

    txt = 'Txt'                                     # Text for default value?
    ftxt = []                                       # Declaration ng list
    print(filename)

    self.filename = filename[0]
    pdf = PdfFileReader(self.filename)
    page_1 = pdf.getPage(0)
    page_1_text = page_1.extractText()
    ftxt = [page_1_text]
    #document = Document(page_1_text)
    #for para in document.paragraphs:
    #    ftxt.append(para.text)
    txt = str(ftxt)

    print('RAW: \n---------------------------------------------------------------------------------------------------')
    print(txt)
    
    stopWords = set(stopwords.words('english'))
    words = word_tokenize(txt)
    freqTable = dict()
    for word in words:
        word = word.lower()
        if word in stopWords:
            continue
        if word in freqTable:
            freqTable[word] +=1
        else:
            freqTable[word] = 1
    sentences = sent_tokenize(txt)
    sentencesValue = dict()
    for sentence in sentences:
            for word, freq in freqTable.items():
                if word in sentence.lower():
                    if sentence in sentencesValue:
                        sentencesValue[sentence] += freq
                    else:
                        sentencesValue[sentence] = freq
    sumValues = 0
    for sentence in sentencesValue:
        sumValues += sentencesValue[sentence]
        
    average = int(sumValues / len(sentencesValue))
    
    summary=''
    for sentence in sentences:
        if(sentence in sentences) and (sentencesValue[sentence] > (.9*average)):
            summary += "" + sentence
            
    print('SUMMARIZE: \n---------------------------------------------------------------------------------------------------')
    
    summary = str(summary).replace("'","")
    summary = str(summary).replace('\\xa0','')
    summary = str(summary).replace('\\n','')
    print(summary)
    mydoc = Document()
    mydoc.add_heading('SUMMARY!')
    mydoc.add_paragraph(summary)
    
    mydoc.save('\summary.docx')
